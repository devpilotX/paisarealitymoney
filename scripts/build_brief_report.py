# Builds PaisaReality_Brief_Report.docx - a short, customer-facing project report
# written in the owner's voice. Run: python scripts/build_brief_report.py
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x00, 0x7A, 0x78)
TEAL_DARK = RGBColor(0x00, 0x53, 0x51)
INK = RGBColor(0x2B, 0x2B, 0x2B)
GRAY = RGBColor(0x5A, 0x5A, 0x5A)

doc = Document()

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK
pf = normal.paragraph_format
pf.line_spacing = 1.18
pf.space_after = Pt(8)

for s in doc.sections:
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)
    s.left_margin = Inches(1.0)
    s.right_margin = Inches(1.0)


def add_bottom_border(paragraph, color="007A78", size="14"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def title_block():
    p = doc.add_paragraph()
    r = p.add_run('Paisa Reality')
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = TEAL
    p.paragraph_format.space_after = Pt(2)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Project Report')
    r2.font.size = Pt(15)
    r2.font.color.rgb = GRAY
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    r3 = p3.add_run("India's one stop money hub")
    r3.font.size = Pt(11)
    r3.font.italic = True
    r3.font.color.rgb = GRAY
    p3.paragraph_format.space_after = Pt(8)
    add_bottom_border(p3)

    meta = doc.add_paragraph()
    mr = meta.add_run('Prepared by Dipanshu Kumar    |    paisareality.com    |    June 2026')
    mr.font.size = Pt(9.5)
    mr.font.color.rgb = GRAY
    meta.paragraph_format.space_after = Pt(14)


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = TEAL
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(12.5)
    r.font.bold = True
    r.font.color.rgb = TEAL_DARK
    return p


def body(text):
    return doc.add_paragraph(text)


def bullet_named(name, rest):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name)
    r.font.bold = True
    p.add_run(rest)
    return p


def shaded_box(lead, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)

    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8F3F2')
    tcPr.append(shd)

    mar = OxmlElement('w:tcMar')
    for side in ('top', 'start', 'bottom', 'end'):
        m = OxmlElement('w:' + side)
        m.set(qn('w:w'), '140')
        m.set(qn('w:type'), 'dxa')
        mar.append(m)
    tcPr.append(mar)

    first = cell.paragraphs[0]
    first.text = ''
    r = first.add_run(lead)
    r.font.bold = True
    r.font.color.rgb = TEAL_DARK
    first.paragraph_format.space_after = Pt(6)

    for ln in lines:
        para = cell.add_paragraph()
        para.add_run(ln)
        para.paragraph_format.space_after = Pt(5)
    return table


# ---------------- content ----------------
title_block()

body(
    "So here is the whole thing, start to finish, in plain words. Paisa Reality is a free "
    "money website I built for normal people in India. Not for the rich, not for experts. "
    "For the person who does not have a personal CA or a wealth manager, who just wants to "
    "know the real gold rate today, or whether they should prepay their loan or invest that "
    "extra money, or which government scheme they actually qualify for. I kept coming back to "
    "one simple feeling. In India, money information is either locked behind a paywall, or it "
    "is buried under so much jargon that a normal person gives up halfway. I really did not "
    "like that. So I wanted to build one honest place where you can check a price, run a real "
    "calculation, find a scheme you qualify for, and see where your own money life actually "
    "stands, all for free and all in simple language. That is the heart of it. This report "
    "walks you through what I have built, what is inside it today, how it works under the "
    "hood, where it stands right now, and what I want to do next. I have kept it short and "
    "honest. No big claims. Just the real picture."
)

h1('The short version')
shaded_box(
    'If you only read one part, read this.',
    [
        'Paisa Reality is a free financial platform for India. It is live at paisareality.com.',
        'It does five big things. Daily prices for gold, silver, petrol, diesel and LPG across '
        'more than 50 cities. Nine Smart Tools that run real simulations, not just simple '
        'formulas. Ten basic calculators like EMI, SIP, FD and income tax. A Money Health Score '
        'that rates your whole money life out of 900. And a finder that matches you to '
        'government schemes you actually qualify for.',
        'On top of that, there are user accounts, a premium plan, an admin dashboard, a full '
        'email system, and a Hindi version of the site.',
        'It is built on a modern, fast and secure setup. It is live, it is tested, and it is '
        'ready for people to use today.',
    ],
)

h1('What Paisa Reality is')
body(
    "The simple way to say it is this. Paisa Reality is India's one stop money hub. One place "
    "for the everyday money questions that keep coming up in a normal Indian household."
)
body(
    "Most sites pick one job. Some only show prices. Some only have calculators. Some only list "
    "schemes. I did not want that. I wanted prices, tools, schemes, bank rates and a personal "
    "money score all sitting together in one clean place. So a person can start with a small "
    "question like what is the gold rate today, and slowly end up actually planning their "
    "retirement or fixing their loans. That is the journey I built it for."
)

h1('What is inside it')
body("Here is everything that is in there today, walked through one by one.")

h2('Daily prices')
body(
    "Every single day, people search for the gold rate, the silver rate, the petrol and diesel "
    "price, the LPG cylinder price. So this is the front door of the site. You get these prices "
    "for more than 50 Indian cities, with clean charts and simple tables. The gold and silver "
    "rates come from a live source and get converted into Indian prices and city rates. This is "
    "the part that brings most people in, because almost everyone checks these numbers."
)

h2('Smart Tools')
body(
    "This is the part I am most proud of, and honestly it is what makes Paisa Reality different. "
    "Most calculator sites just run one textbook formula and hand you a number. These do not. "
    "There are nine Smart Tools, and they do the real, hard math."
)
bullet_named('Retirement Optimizer. ', 'It runs 10,000 different market scenarios to tell you how much you need to retire and what monthly SIP gets you there.')
bullet_named('Prepay vs Invest. ', 'Should you prepay your home loan or invest the money? It gives a risk adjusted, after tax answer, with the odds of each path winning.')
bullet_named('Debt Optimizer. ', 'Got many loans? It finds the cheapest and fastest order to clear them, and it understands the tax angle too.')
bullet_named('Lifecycle Tax Optimizer. ', 'Old regime or new regime, checked across your whole career instead of just one single year.')
bullet_named('Budget Optimizer. ', 'It goes past the old 50/30/20 rule, finds your real surplus, and flags where the money is leaking.')
bullet_named('Tax Harvesting. ', 'Which holdings to sell before year end to cut your capital gains tax, using the 1.25 lakh exemption.')
bullet_named('Gold Planner. ', 'Gold returns and risk, and whether to buy in one shot or spread it out over time.')
bullet_named('Scheme Maximizer. ', 'The total rupee benefit of every central scheme you qualify for, with the steps to claim each one.')
bullet_named('Salary Optimizer. ', 'The best CTC breakup to legally bring your income tax down.')
body(
    "And here is the nice part. All of this runs inside your own browser. Your numbers never "
    "leave your device. So it is private by design."
)

h2('Basic calculators')
body(
    "Not everyone needs the heavy tools. Sometimes you just want a quick number. So there are "
    "ten simple calculators too. EMI, SIP, FD, PPF, income tax, home loan, NPS, gratuity, HRA "
    "and inflation. Clean, fast, no sign up needed. Just the everyday stuff."
)

h2('Money Health Score')
body(
    "This one is special. The Money Health Score gives you a single number out of 900 for your "
    "entire money life. Think of it like a CIBIL score, but for everything, not just your loans."
)
body(
    "It looks at eight parts of your money. Your savings rate, your emergency fund, your debt, "
    "your retirement, your investing, your insurance, your tax efficiency, and your day to day "
    "money habits. It puts them together into one score and tells you the band you are in, from "
    "At Risk at the bottom to Excellent at the top. Then it points you to the exact tool to fix "
    "your weakest area. You can save your score, track it over time, and share it. It turns a "
    "vague feeling of am I doing okay with money into one clear number you can actually act on."
)

h2('Government schemes')
body(
    "India has hundreds of government schemes, and the real problem is not the schemes, it is "
    "finding the ones that fit you. So I built a matcher. You fill a short form about yourself, "
    "and it shows you the central and state schemes you actually qualify for, instead of making "
    "you read through hundreds of pages. The Scheme Maximizer in the Smart Tools goes one step "
    "further and adds up the real rupee benefit."
)

h2('Bank rates')
body(
    "People also want to compare banks before they park their money or take a loan. So the site "
    "compares FD rates, savings rates, home loan rates and personal loan rates across more than "
    "50 banks, in simple tables. Quick to scan, easy to compare."
)

h2('Accounts, premium, admin and email')
body(
    "Behind all of this there is a full system. People can create an account, verify their "
    "email, log in safely, and reset a forgotten password. There is a premium plan handled "
    "through Razorpay for the people who want more. There is an admin dashboard where I manage "
    "blog posts, prices, user messages and emails from one place. And there is a proper email "
    "system for welcome mails, verification, password resets and newsletters."
)

h2('Two languages')
body(
    "India is not an English only country, and I did not want to forget that. So there is a "
    "Hindi version of the site as well, so more people can use it in the language they actually "
    "think in."
)

h1('How it is built')
body(
    "I will keep this part simple, even though it is the technical bit. The site is built on "
    "Next.js with React, which is a modern and fast setup that a lot of big companies use. The "
    "code is written in TypeScript in strict mode, which basically means the code checks itself "
    "for mistakes before anything goes live. The design is done with Tailwind, so it stays clean "
    "and works on a phone and a laptop alike."
)
body(
    "All the data, the prices, the schemes, the banks, the users and the scores, sits in a "
    "PostgreSQL database, which is a solid and trusted choice. Logins are protected properly, "
    "passwords are hashed with bcrypt, and sessions use secure tokens. Payments go through "
    "Razorpay. Emails go through Resend. Reports can be made as PDF files. There are also safety "
    "layers like rate limiting and input cleaning, so the site does not get abused. And like I "
    "said, the Smart Tools run in your browser, so your private money numbers stay with you."
)

h1('Where it stands today')
body(
    "Here is the honest status. The site is live at paisareality.com and working. The code "
    "builds cleanly and passes its strict checks. The heavy money tools, the ones doing the real "
    "math, have automated tests around them, so I can change things without quietly breaking the "
    "numbers. The SEO work is done too, the sitemap, the structured data, the FAQ sections and "
    "the per page details that help Google show the site to the right people. In short, this is "
    "not a demo. It is a real, working product that people can use today."
)

h1('A few honest notes')
body(
    "A couple of honest things, because I would rather tell you straight. These tools are made "
    "to teach and to estimate. They are not financial advice, and the site says that clearly. "
    "For a big life decision, you should still talk to a proper advisor. And Paisa Reality is "
    "built and run by me, so it grows piece by piece. Some parts are deeper than others today, "
    "and I keep improving them. I would rather build it honestly and steadily than fake a "
    "finished look."
)

h1('What is next')
body(
    "Where I want to take it from here. I want to make the daily prices even more automatic and "
    "cover more cities. I want to add more Smart Tools, because that is the real edge of this "
    "site. I want a deeper, more guided Money Health Score that walks a person through fixing "
    "each weak spot one by one. More content and more Hindi, so it reaches more people. And "
    "slowly, a gentle premium plan for the people who want the advanced reports, while the core "
    "stays free for everyone. The big goal does not change. One honest, simple, free place for an "
    "ordinary Indian to deal with money."
)

h1('Closing')
body(
    "That is Paisa Reality, start to finish. I did not build it to look flashy. I built it "
    "because I genuinely wanted a place like this to exist. A place that treats a normal "
    "person's money questions with respect and gives real answers for free. It is live, it is "
    "real, and it is only going to get better. Thank you for reading, and thank you for giving "
    "it a look."
)

out = 'PaisaReality_Brief_Report.docx'
doc.save(out)
print('saved', out)
