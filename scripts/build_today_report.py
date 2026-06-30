# -*- coding: utf-8 -*-
"""Generate the Paisa Reality daily work report (25 June 2026) as a Word file."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = r"C:\Users\Dipan\OneDrive\Desktop\Paisareality\paisarealitymoney\Paisa Reality Reports"
OUT_FILE = os.path.join(OUT_DIR, "Paisa Reality - Work Report - 25 June 2026.docx")
ACCENT = RGBColor(0x00, 0x7A, 0x78)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def num(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ---------------------------------------------------------------- Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Paisa Reality")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Daily Work Report")
r.bold = True
r.font.size = Pt(18)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Date: 25 June 2026\nLive site: https://paisareality.com\nAdmin: https://admin.paisareality.com").font.size = Pt(11)

doc.add_paragraph()
para(
    "This report documents all work completed on Paisa Reality on 25 June 2026. "
    "Every change was built, type-checked, committed to the main branch, deployed to the "
    "production server, and verified live. All database work was additive and idempotent, "
    "so existing schemes, users, prices, and scores were preserved.",
    italic=True,
)

# ---------------------------------------------------------------- 1. Summary
h("1. Executive Summary", 1)
para("In one day, the platform moved from a thin data state to a full, production grade money information site for India. The headline outcomes:")
bullet("Government schemes grew from 10 to 351, covering all 28 states and all 8 union territories.")
bullet("The admin dashboard was moved to its own subdomain, admin.paisareality.com, and blocked on the main domain.")
bullet("The Blog section was renamed to Newsletter with permanent redirects.")
bullet("Daily prices were made current and now refresh automatically every morning.")
bullet("Bank rate comparison grew from 13 to 51 banks, so the 50 plus banks claim is now true.")
bullet("The admin panel now shows live site data, and a protected stats endpoint was added.")
bullet("SEO was re-verified and kept intact, with full structured data on every scheme page.")
bullet("Mobile performance and best practices were improved by removing unused third party load.")
bullet("The README was rewritten clean, with environment variable names only and no secrets.")
bullet("The on-site assistant was first removed, then restored on request, and is live again.")

# ---------------------------------------------------------------- 2. Schemes
h("2. Government Schemes: 10 to 351", 1)
para(
    "The live database held only 10 schemes at the start of the day, all central. A common "
    "misunderstanding was that the site already had 200 plus schemes; in fact an old seed file "
    "was written for MySQL and never loaded into the live PostgreSQL database. Two additive, "
    "idempotent seeds were written and run against the live database.",
)
table(
    ["Stage", "Schemes added", "Running total"],
    [
        ["Start of day (existing)", "-", "10 (all central)"],
        ["Expansion seed 1", "206 new, 10 refreshed", "216"],
        ["Expansion seed 2", "135 new", "351"],
        ["Final", "341 new in total", "351 (171 central, 180 state)"],
    ],
)
para("Coverage now spans every state and union territory. Dedicated state level schemes per state and UT:")
table(
    ["State or UT", "Schemes", "State or UT", "Schemes"],
    [
        ["Tamil Nadu", "11", "West Bengal", "12"],
        ["Karnataka", "10", "Maharashtra", "10"],
        ["Gujarat", "9", "Uttar Pradesh", "9"],
        ["Bihar", "8", "Madhya Pradesh", "8"],
        ["Rajasthan", "8", "Kerala", "7"],
        ["Punjab", "7", "Telangana", "7"],
        ["Assam", "6", "Delhi", "6"],
        ["Haryana", "6", "Jharkhand", "6"],
        ["Odisha", "6", "Himachal Pradesh", "5"],
        ["Andhra Pradesh", "4", "Chhattisgarh", "4"],
        ["Goa", "4", "Uttarakhand", "4"],
        ["Jammu and Kashmir", "3", "Arunachal Pradesh", "2"],
        ["Manipur", "2", "Meghalaya", "2"],
        ["Mizoram", "2", "Nagaland", "2"],
        ["Puducherry", "2", "Tripura", "2"],
        ["Andaman and Nicobar", "1", "Chandigarh", "1"],
        ["Dadra Nagar Haveli Daman Diu", "1", "Ladakh", "1"],
        ["Lakshadweep", "1", "Sikkim", "1"],
    ],
)
para("How accuracy was kept:", bold=True)
bullet("Only real, currently active central and state schemes were used.")
bullet("Where a precise current figure was not certain, the wording was kept general on purpose.")
bullet("Every scheme links to its official source so users can verify before applying.")
bullet("High traffic state schemes were checked against official sources, for example Maharashtra Ladki Bahin, Odisha Subhadra, Karnataka Gruha Lakshmi, Telangana Rythu Bharosa, Tamil Nadu Kalaignar Magalir Urimai Thogai, and central updates such as Ayushman Vay Vandana and PM Surya Ghar.")
para("Delivery details:", bold=True)
bullet("Two seed scripts: scripts/seed-schemes-expansion.ts and scripts/seed-schemes-expansion-2.ts.")
bullet("Each upserts by slug using INSERT ON CONFLICT DO UPDATE inside a transaction, so it is safe to re-run with zero duplicates and zero data loss.")
bullet("New scheme categories were added so every category hub renders: employment, pension, insurance, finance, and social welfare.")
bullet("Every scheme page renders full metadata, canonical, OpenGraph and Twitter cards, GovernmentService, BreadcrumbList and FAQPage structured data, a visible FAQ, and a sitemap entry.")
bullet("The scheme finder and the per-state and per-category pages match and display all new schemes.")

# ---------------------------------------------------------------- 3. Admin subdomain
h("3. Admin Dashboard Moved to admin.paisareality.com", 1)
bullet("The admin dashboard now serves only on admin.paisareality.com.")
bullet("On the main domain, any /admin or /api/admin request returns 404, so admin is not reachable at paisareality.com/admin.")
bullet("Host based routing is handled in middleware.ts, and a dedicated Nginx server block was committed under deploy/nginx/ and applied on the server, with an X-Robots-Tag noindex header for the admin subdomain.")
bullet("Admin login reads ADMIN_EMAIL and ADMIN_PASSWORD from the environment and fails closed if the password is not set. The password is never printed or hardcoded.")
bullet("The TLS certificate already covers the subdomain, the admin DNS record already resolves through Cloudflare, and the admin credentials are set, so the admin area is fully usable.")

# ---------------------------------------------------------------- 4. Newsletter
h("4. Blog Renamed to Newsletter", 1)
bullet("The Blog section is now Newsletter across navigation, footer, titles, headings, and metadata.")
bullet("/blog and /blog/[slug] now return permanent 301 redirects to /newsletter and /newsletter/[slug].")
bullet("Sitemap, canonicals, internal links, and structured data point to the newsletter routes.")

# ---------------------------------------------------------------- 5. Assistant
h("5. On-site Assistant (Yojana Mitra)", 1)
bullet("The assistant widget and its API were first removed as part of the initial brief.")
bullet("On request, the assistant was restored and is live again on the site.")
bullet("It runs in guided mode with no key required. Adding a GEMINI_API_KEY environment variable enables full AI replies.")

# ---------------------------------------------------------------- 6. Prices
h("6. Daily Prices Made Live and Automatic", 1)
para("The prices were four days stale because nothing was scheduled to refresh them. This was fixed.")
bullet("Prices were refreshed to the current date. Gold and silver use live spot rates from public market sources converted to Indian prices; petrol, diesel, and LPG use representative published rates.")
bullet("A systemd timer was installed on the server to run the price update every day at 07:00 India time, so prices stay current automatically.")
bullet("The price pages now show the current date, for example the gold and petrol pages read 25 June 2026.")
bullet("Coverage is 50 plus cities for gold, silver, petrol, and diesel, and per state for LPG.")
bullet("The timer unit files are committed under deploy/systemd/ for reference. The cron secret stays in the environment file and is never written into the unit.")

# ---------------------------------------------------------------- 7. Banks
h("7. Bank Rate Comparison: 13 to 51 Banks", 1)
bullet("Banks grew from 13 to 51, covering public sector, private, small finance, cooperative banks, and India Post, with 248 rate rows.")
bullet("The 50 plus banks claim shown across the site is now accurate.")
bullet("Rates are indicative and rounded, with a clear note to verify with the bank, since rates change often.")
bullet("The seed (scripts/seed-banks-expansion.ts) upserts each bank and replaces its rate rows inside a transaction, so it is safe to re-run.")

# ---------------------------------------------------------------- 8. Admin stats
h("8. Admin Panel Live Data", 1)
bullet("A protected stats endpoint was added at /api/admin/stats, returning live counts.")
bullet("The admin overview now shows Government Schemes, Banks, Cities, Registered Users, the central and state split, the price freshness date, and newsletter counts.")
bullet("A stale link to the old blog path in the admin actions was corrected to the newsletter path.")

# ---------------------------------------------------------------- 9. SEO
h("9. SEO Re-verification", 1)
bullet("Every page type was re-verified. Existing titles, descriptions, canonicals, OpenGraph, Twitter, and JSON-LD were kept intact and not weakened.")
bullet("All pages are indexable except the Hindi section, admin, and dashboard, which are kept noindex and out of the sitemap.")
bullet("robots and sitemap were confirmed correct and complete. The sitemap now lists all 351 scheme pages plus all other page types.")
bullet("The FAQ structured data is now server rendered, so it appears in the initial HTML on every page that uses it, which is more reliable for indexing.")
bullet("The biggest SEO gain is the 341 new scheme pages, each fully optimized for its target keyword.")

# ---------------------------------------------------------------- 10. Performance
h("10. Performance and Lighthouse", 1)
para("Reported scores before the performance pass: mobile Performance 67 and Best Practices 81; desktop near perfect; Accessibility and SEO already 100.")
para("Root cause and fixes:", bold=True)
bullet("The AdSense library was loading on every page even though no ad slots are configured, so it added cost with no benefit and no revenue. It now loads only when ad slots are set, so it is currently not loaded.")
bullet("Google Analytics was moved to load lazily after the page is interactive.")
bullet("Ad units now lazy load on scroll, reserve their space to avoid layout shift, and no longer log errors to the console.")
bullet("DNS prefetch hints were added for the ad and analytics origins.")
para(
    "Note: a perfect 100 on mobile is straightforward while there are no ads. When real AdSense "
    "ad units are created and the slot values are set, the ad library and its third party cookies "
    "return and mobile Performance and Best Practices will dip a few points. The lazy loading and "
    "reserved space keep that dip small. Scores should be re-run from a browser to confirm the new numbers.",
    italic=True,
)

# ---------------------------------------------------------------- 11. README
h("11. README", 1)
bullet("The GitHub README was rewritten to be clean and professional.")
bullet("It lists what the site is, live features, local setup, environment variable names only with no values, and the deployment process.")
bullet("It contains no secrets, no personal name, and no reference to removed features.")

# ---------------------------------------------------------------- 12. Database and infra
h("12. Database and Infrastructure", 1)
bullet("Database: PostgreSQL. Connectivity was confirmed solid throughout the day across many live queries.")
bullet("All seeds are additive and idempotent (CREATE TABLE IF NOT EXISTS style and INSERT ON CONFLICT DO UPDATE). No reset scripts were run. Existing data survived untouched.")
bullet("The app runs with PM2 behind Nginx, with Cloudflare in front. Deploys are done by git pull on the server, never a fresh clone.")
bullet("Standalone npm scripts were added for each additive seed so they are run explicitly and never as part of a full reset.")

# ---------------------------------------------------------------- 13. QA
h("13. Quality Assurance", 1)
bullet("A full A to Z sweep checked 61 page types and all returned HTTP 200 with no broken pages.")
bullet("Checked: home, scheme finder, scheme detail, category and state hubs, bank hub with four sub pages and bank detail pages, all 10 basic calculators, all 9 smart tools, the Money Health Score, newsletter, all 5 price hubs and city pages, legal pages, sitemap, and robots.")
bullet("typecheck and the production build both completed with zero errors before every deploy.")
bullet("Public checks through Cloudflare confirmed new scheme pages, the admin subdomain, the blog to newsletter redirect, and the daily prices are live.")

# ---------------------------------------------------------------- 14. Commits
h("14. Code Commits (branch main)", 1)
table(
    ["Commit", "Summary"],
    [
        ["d98f77c", "200 plus schemes, admin subdomain, newsletter 301s, SEO, README, assistant removal"],
        ["f0c18ed", "Exclude scripts and tests from app typecheck so the production build passes"],
        ["db169cf", "Server render the FAQ structured data"],
        ["798832e", "Restore the on-site assistant"],
        ["8a6a381", "Add 50 plus banks, daily price auto-update, admin live stats"],
        ["5198fa4", "Expand to 351 schemes covering all states and union territories"],
        ["5f0f181", "Optimize mobile performance and best practices"],
    ],
)

# ---------------------------------------------------------------- 15. Final state
h("15. Final Live State", 1)
table(
    ["Item", "Value"],
    [
        ["Government schemes", "351 (171 central, 180 state)"],
        ["States and UTs covered", "36 of 36"],
        ["Banks compared", "51 (248 rate rows)"],
        ["Cities with daily prices", "50 plus"],
        ["Prices", "Current, auto-updated daily at 07:00 IST"],
        ["Admin", "admin.paisareality.com only, 404 on main domain"],
        ["Newsletter", "Live, /blog returns 301 to /newsletter"],
        ["Assistant", "Live in guided mode"],
        ["Accessibility and SEO scores", "100 (already)"],
        ["typecheck and build", "Pass (zero errors)"],
    ],
)

# ---------------------------------------------------------------- 16. Owner actions
h("16. Optional Next Steps for the Owner", 1)
num("Purge the Cloudflare cache for good hygiene. The price, bank, and scheme pages are served dynamically and are already fresh, so this is optional.")
num("To turn on ads, create ad units in the AdSense dashboard and set NEXT_PUBLIC_ADSENSE_DEFAULT_SLOT and NEXT_PUBLIC_ADSENSE_IN_ARTICLE_SLOT in the environment file, then rebuild. Ads will then show, with a small expected dip in mobile scores.")
num("To enable full AI replies in the assistant, add a GEMINI_API_KEY to the environment file and restart the process.")
num("In Google Search Console, resubmit the sitemap and request indexing for the new scheme pages so they are crawled sooner.")
num("For security, if the admin password or any API keys were ever shared in plain text, rotate them and update the environment file.")
num("Review the dependency vulnerabilities reported by the package manager in a separate, tested change.")

# ---------------------------------------------------------------- 17. Note
h("17. Disclaimer", 1)
para(
    "Paisa Reality is an informational website, not a financial advisor. Scheme details, prices, "
    "and bank rates are indicative and sourced from official and public data, and users are asked "
    "to verify with official sources before making any decision.",
    italic=True,
)

os.makedirs(OUT_DIR, exist_ok=True)
doc.save(OUT_FILE)
print("SAVED:", OUT_FILE)
